import os
import pandas as pd
from tkinter import Tk, Label, Button, Entry, filedialog, messagebox, OptionMenu, StringVar, Toplevel, ttk, Checkbutton, BooleanVar, Canvas, Scrollbar, Frame
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from tkinter.font import Font
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Initialize the main application window
root = Tk()
root.title("Certificate and Email Sender")
root.geometry("600x600")
root.configure(bg="#f0f4f8")

# Define global variables
excel_path = None
template_path = None
email_template_path = None
output_format = StringVar(value="PDF")
selected_column = StringVar(value="Select Column")
email_column = StringVar(value="Select Column")
subject_column = StringVar(value="Select Column")
content_column = StringVar(value="Select Column")
column_options = ["Select Column"]
placeholder_map = {}
email_placeholder_map = {}
send_certificates = BooleanVar(value=True)

# Define dropdown menu variables globally
column_menu = None
email_column_menu = None
subject_column_menu = None
content_column_menu = None

# Custom font
custom_font = Font(family="Comic Sans MS", size=12)

# SMTP Configuration
smtp_server = StringVar()
smtp_port = StringVar()
smtp_user = StringVar()
smtp_password = StringVar()

def select_excel_file():
    global excel_path, column_options
    excel_path = filedialog.askopenfilename(
        title="Select Excel File", filetypes=(("Excel files", "*.xlsx"), ("All files", "*.*"))
    )
    if excel_path:
        excel_label.config(text=f"Selected: {os.path.basename(excel_path)}")
        df = pd.read_excel(excel_path)
        global column_options
        column_options = df.columns.tolist()
        update_option_menus()

def select_template_file():
    global template_path
    template_path = filedialog.askopenfilename(
        title="Select Template File", filetypes=(("Word files", "*.docx"), ("All files", "*.*"))
    )
    if template_path:
        template_label.config(text=f"Selected: {os.path.basename(template_path)}")
        update_placeholder_mapping()

def select_email_template_file():
    global email_template_path
    email_template_path = filedialog.askopenfilename(
        title="Select Email Template File", filetypes=(("Word files", "*.docx"), ("All files", "*.*"))
    )
    if email_template_path:
        email_template_label.config(text=f"Selected: {os.path.basename(email_template_path)}")
        update_email_placeholder_mapping()

def update_placeholder_mapping():
    global placeholder_map
    if not template_path:
        return

    placeholder_map = {}
    doc = Document(template_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" in run.text and "}}" in run.text:
                placeholder_text = run.text.split("{{")[1].split("}}")[0].strip()
                placeholders.add(placeholder_text)

    mapping_window = Toplevel(root)
    mapping_window.title("Map Certificate Placeholders")
    mapping_window.geometry("400x400")
    mapping_window.configure(bg="#f0f4f8")

    for i, placeholder in enumerate(placeholders):
        Label(mapping_window, text=f"Placeholder: {{ {placeholder} }}", bg="#f0f4f8").grid(row=i, column=0, padx=10, pady=5)
        column_var = StringVar(mapping_window)
        column_var.set("Select Column")
        placeholder_map[placeholder] = column_var
        OptionMenu(mapping_window, column_var, *column_options).grid(row=i, column=1, padx=10, pady=5)

    Button(mapping_window, text="Done", command=mapping_window.destroy, bg="#4CAF50", fg="white").grid(row=len(placeholders), column=0, columnspan=2, pady=20)

def update_email_placeholder_mapping():
    global email_placeholder_map
    if not email_template_path:
        return

    email_placeholder_map = {}
    doc = Document(email_template_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" in run.text and "}}" in run.text:
                placeholder_text = run.text.split("{{")[1].split("}}")[0].strip()
                placeholders.add(placeholder_text)

    mapping_window = Toplevel(root)
    mapping_window.title("Map Email Placeholders")
    mapping_window.geometry("400x400")
    mapping_window.configure(bg="#f0f4f8")

    for i, placeholder in enumerate(placeholders):
        Label(mapping_window, text=f"Placeholder: {{ {placeholder} }}", bg="#f0f4f8").grid(row=i, column=0, padx=10, pady=5)
        column_var = StringVar(mapping_window)
        column_var.set("Select Column")
        email_placeholder_map[placeholder] = column_var
        OptionMenu(mapping_window, column_var, *column_options).grid(row=i, column=1, padx=10, pady=5)

    Button(mapping_window, text="Done", command=mapping_window.destroy, bg="#4CAF50", fg="white").grid(row=len(placeholders), column=0, columnspan=2, pady=20)

def update_option_menus():
    updated_options = ["Select Column"] + column_options
    selected_column.set(updated_options[0])
    email_column.set(updated_options[0])
    subject_column.set(updated_options[0])
    content_column.set(updated_options[0])
    
    # Update the options for each dropdown menu
    if column_menu:
        column_menu["menu"].delete(0, "end")
        for option in updated_options:
            column_menu["menu"].add_command(label=option, command=lambda value=option: selected_column.set(value))
    
    if email_column_menu:
        email_column_menu["menu"].delete(0, "end")
        for option in updated_options:
            email_column_menu["menu"].add_command(label=option, command=lambda value=option: email_column.set(value))
    
    if subject_column_menu:
        subject_column_menu["menu"].delete(0, "end")
        for option in updated_options:
            subject_column_menu["menu"].add_command(label=option, command=lambda value=option: subject_column.set(value))
    
    if content_column_menu:
        content_column_menu["menu"].delete(0, "end")
        for option in updated_options:
            content_column_menu["menu"].add_command(label=option, command=lambda value=option: content_column.set(value))

def generate_and_send_certificates():
    if not excel_path or not template_path or selected_column.get() == "Select Column" or email_column.get() == "Select Column":
        messagebox.showerror("Error", "Please select Excel file, Template file, Name column, and Email column.")
        return

    df = pd.read_excel(excel_path)
    font_name = 'Comic Sans MS'
    font_size = Pt(16)
    output_dir = "certificates/"
    os.makedirs(output_dir, exist_ok=True)

    name_counter = {}
    progress = ttk.Progressbar(root, length=400, mode='determinate')
    progress.pack(pady=10)

    total = len(df)
    for index, row in df.iterrows():
        if send_certificates.get() and template_path:
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    for placeholder, column_var in placeholder_map.items():
                        column = column_var.get()
                        if column != "Select Column":
                            placeholder_format = f"{{{{{placeholder}}}}}"
                            if placeholder_format in run.text:
                                run.text = run.text.replace(placeholder_format, str(row[column]))
                                run.bold = True
                                run.italic = True
                                run.font.name = font_name
                                run.font.size = font_size

            name_value = str(row[selected_column.get()])
            if name_value in name_counter:
                name_counter[name_value] += 1
                unique_name = f"{name_value}_{name_counter[name_value]}"
            else:
                name_counter[name_value] = 1
                unique_name = name_value

            if output_format.get() == "PDF":
                docx_path = os.path.join(output_dir, f'{unique_name}_certificate.docx')
                pdf_path = os.path.join(output_dir, f'{unique_name}_certificate.pdf')
                doc.save(docx_path)
                convert(docx_path, pdf_path)
                os.remove(docx_path)
                file_path = pdf_path
            else:
                file_path = os.path.join(output_dir, f'{unique_name}_certificate.docx')
                doc.save(file_path)

            if email_template_path:
                email_subject = row.get(subject_column.get(), "")
                email_body = row.get(content_column.get(), "")
                send_email(row[email_column.get()], email_subject, email_body, file_path)

        progress['value'] = (index + 1) / total * 100
        root.update_idletasks()

    progress.pack_forget()
    messagebox.showinfo("Success", "Certificates generated and emails sent.")

def send_email(recipient, subject, body, attachment_path=None):
    try:
        msg = MIMEMultipart()
        msg['From'] = smtp_user.get()
        msg['To'] = recipient
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        if attachment_path:
            with open(attachment_path, "rb") as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    "Content-Disposition",
                    f"attachment; filename= {os.path.basename(attachment_path)}",
                )
                msg.attach(part)

        with smtplib.SMTP(smtp_server.get(), int(smtp_port.get())) as server:
            server.starttls()
            server.login(smtp_user.get(), smtp_password.get())
            server.sendmail(msg['From'], msg['To'], msg.as_string())
        print(f"Email sent to {recipient}")
    except Exception as e:
        print(f"Failed to send email: {e}")

def send_email_only():
    if not excel_path or email_column.get() == "Select Column":
        messagebox.showerror("Error", "Please select Excel file and Email column.")
        return

    df = pd.read_excel(excel_path)
    progress = ttk.Progressbar(root, length=400, mode='determinate')
    progress.pack(pady=10)

    total = len(df)
    for index, row in df.iterrows():
        send_email(row[email_column.get()], subject=row.get(subject_column.get(), ""), body=row.get(content_column.get(), ""))

        progress['value'] = (index + 1) / total * 100
        root.update_idletasks()

    progress.pack_forget()
    messagebox.showinfo("Success", "Emails sent.")

def create_widgets():
    global column_menu, email_column_menu, subject_column_menu, content_column_menu

    # Create scrollable frame
    canvas = Canvas(root)
    scrollbar = Scrollbar(root, orient="vertical", command=canvas.yview)
    scrollable_frame = Frame(canvas)
    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Widgets
    Label(scrollable_frame, text="SMTP Server:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    Entry(scrollable_frame, textvariable=smtp_server, bg="#ffffff", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="SMTP Port:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    Entry(scrollable_frame, textvariable=smtp_port, bg="#ffffff", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="SMTP User:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    Entry(scrollable_frame, textvariable=smtp_user, bg="#ffffff", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="SMTP Password:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    Entry(scrollable_frame, textvariable=smtp_password, show="*", bg="#ffffff", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="Excel File:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    global excel_label
    excel_label = Label(scrollable_frame, text="No Excel file selected", bg="#f0f4f8", font=custom_font)
    excel_label.pack(pady=5)
    Button(scrollable_frame, text="Browse", command=select_excel_file, bg="#4CAF50", fg="white", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="Template File:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    global template_label
    template_label = Label(scrollable_frame, text="No Template file selected", bg="#f0f4f8", font=custom_font)
    template_label.pack(pady=5)
    Button(scrollable_frame, text="Browse", command=select_template_file, bg="#4CAF50", fg="white", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="Email Template File:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    global email_template_label
    email_template_label = Label(scrollable_frame, text="No Email Template file selected", bg="#f0f4f8", font=custom_font)
    email_template_label.pack(pady=5)
    Button(scrollable_frame, text="Browse", command=select_email_template_file, bg="#4CAF50", fg="white", font=custom_font).pack(pady=5)

    Label(scrollable_frame, text="Output Format:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    format_menu = OptionMenu(scrollable_frame, output_format, "PDF", "DOCX")
    format_menu.pack(pady=5)

    Label(scrollable_frame, text="Select Column for Name:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    column_menu = OptionMenu(scrollable_frame, selected_column, *column_options)
    column_menu.pack(pady=5)

    Label(scrollable_frame, text="Select Email Column:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    email_column_menu = OptionMenu(scrollable_frame, email_column, *column_options)
    email_column_menu.pack(pady=5)

    Label(scrollable_frame, text="Select Subject Column:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    subject_column_menu = OptionMenu(scrollable_frame, subject_column, *column_options)
    subject_column_menu.pack(pady=5)

    Label(scrollable_frame, text="Select Content Column:", bg="#f0f4f8", font=custom_font).pack(pady=5)
    content_column_menu = OptionMenu(scrollable_frame, content_column, *column_options)
    content_column_menu.pack(pady=5)

    Checkbutton(scrollable_frame, text="Send with Certificate", variable=send_certificates, bg="#f0f4f8", font=custom_font).pack(pady=5)

    Button(scrollable_frame, text="Generate Certificates and Send Emails", command=generate_and_send_certificates, bg="#4CAF50", fg="white", font=custom_font).pack(pady=10)
    Button(scrollable_frame, text="Send Emails Only", command=send_email_only, bg="#4CAF50", fg="white", font=custom_font).pack(pady=10)

create_widgets()
root.mainloop()
