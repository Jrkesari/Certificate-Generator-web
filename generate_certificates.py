import os
import pandas as pd
from tkinter import Tk, Label, Button, filedialog, messagebox, OptionMenu, StringVar, Toplevel, ttk
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from tkinter.font import Font

# Create the main application window
root = Tk()
root.title("Certificate Generator")
root.geometry("800x700")
root.configure(bg="#f0f4f8")

# Custom font
custom_font = Font(family="Comic Sans MS", size=12)

# Initialize variables
excel_path = None
template_path = None
output_format = StringVar(value="PDF")
selected_column = StringVar(value="Select Column")
column_options = ["Select Column"]
placeholder_map = {}

# Function to select an Excel file
def select_excel_file():
    global excel_path, column_options
    excel_path = filedialog.askopenfilename(
        title="Select Excel File", filetypes=(("Excel files", "*.xlsx"), ("All files", "*.*"))
    )
    if excel_path:
        excel_label.config(text=f"Selected: {os.path.basename(excel_path)}")
        df = pd.read_excel(excel_path)
        column_options = df.columns.tolist()
        selected_column.set(column_options[0] if column_options else "Select Column")
        
        # Update the OptionMenu with new columns
        column_menu['menu'].delete(0, 'end')
        for col in column_options:
            column_menu['menu'].add_command(label=col, command=lambda value=col: selected_column.set(value))

# Function to select a Word template
def select_template_file():
    global template_path
    template_path = filedialog.askopenfilename(
        title="Select Template File", filetypes=(("Word files", "*.docx"), ("All files", "*.*"))
    )
    if template_path:
        template_label.config(text=f"Selected: {os.path.basename(template_path)}")
        update_placeholder_mapping()

# Function to update placeholder-column mapping
def update_placeholder_mapping():
    global placeholder_map
    if not template_path:
        return

    placeholder_map = {}
    # Extract placeholders from the template
    doc = Document(template_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" in run.text and "}}" in run.text:
                placeholder_text = run.text.split("{{")[1].split("}}")[0].strip()
                placeholders.add(placeholder_text)

    # Create a new window for mapping
    mapping_window = Toplevel(root)
    mapping_window.title("Map Placeholders")
    mapping_window.geometry("400x400")
    mapping_window.configure(bg="#f0f4f8")

    for i, placeholder in enumerate(placeholders):
        Label(mapping_window, text=f"Placeholder: {{ {placeholder} }}", bg="#f0f4f8").grid(row=i, column=0, padx=10, pady=5)
        column_var = StringVar(mapping_window)
        column_var.set("Select Column")
        placeholder_map[placeholder] = column_var
        OptionMenu(mapping_window, column_var, *column_options).grid(row=i, column=1, padx=10, pady=5)

    Button(mapping_window, text="Done", command=mapping_window.destroy, bg="#4CAF50", fg="white").grid(row=len(placeholders), column=0, columnspan=2, pady=20)

# Function to generate certificates
def generate_certificates():
    if not excel_path or not template_path or selected_column.get() == "Select Column":
        messagebox.showerror("Error", "Please select Excel file, Template file, and Name column.")
        return

    df = pd.read_excel(excel_path)
    font_name = 'Comic Sans MS'
    font_size = Pt(16)
    output_dir = "certificates/"
    os.makedirs(output_dir, exist_ok=True)

    # Dictionary to track name occurrences
    name_counter = {}

    progress = ttk.Progressbar(root, length=400, mode='determinate')
    progress.pack(pady=10)

    total = len(df)
    for index, row in df.iterrows():
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for placeholder, column_var in placeholder_map.items():
                    column = column_var.get()
                    if column != "Select Column":
                        placeholder_format = f"{{{{{placeholder}}}}}"
                        if placeholder_format in run.text:
                            run.text = run.text.replace(placeholder_format, str(row[column]))
                            run.bold = True
                            run.italic = True
                            run.font.name = font_name
                            run.font.size = font_size

        # Handle duplicate names by appending a counter
        name_value = str(row[selected_column.get()])
        if name_value in name_counter:
            name_counter[name_value] += 1
            unique_name = f"{name_value}_{name_counter[name_value]}"
        else:
            name_counter[name_value] = 1
            unique_name = name_value

        if output_format.get() == "PDF":
            docx_path = os.path.join(output_dir, f'{unique_name}_certificate.docx')
            pdf_path = os.path.join(output_dir, f'{unique_name}_certificate.pdf')
            doc.save(docx_path)
            convert(docx_path, pdf_path)
            os.remove(docx_path)  # Remove DOCX after conversion to PDF
        else:
            docx_path = os.path.join(output_dir, f'{unique_name}_certificate.docx')
            doc.save(docx_path)

        progress['value'] = ((index + 1) / total) * 100
        root.update_idletasks()

    progress.pack_forget()
    messagebox.showinfo("Success", "Certificates generated successfully!")

# Information section
info_label = Label(root, text="How to add placeholders in your DOCX template:", font=("Arial", 14, "bold"), bg="#4CAF50", fg="white", pady=10)
info_label.pack(fill="x")

instructions = """
1. Open your Word document (.docx) file.
2. Add placeholders in the format: {{ placeholder_name }}.
3. For example, use {{ Name }} or {{ Date }} with brackets.
4. Save the document and use it as a template.
"""
instructions_label = Label(root, text=instructions, font=custom_font, bg="#f0f4f8", justify="left", padx=10, pady=10)
instructions_label.pack(fill="x")

# GUI Elements
excel_label = Label(root, text="No Excel file selected", bg="#f0f4f8", font=custom_font)
excel_label.pack(pady=10)

template_label = Label(root, text="No Template file selected", bg="#f0f4f8", font=custom_font)
template_label.pack(pady=10)

select_excel_btn = Button(root, text="Select Excel File", command=select_excel_file, bg="#4CAF50", fg="white", font=custom_font)
select_excel_btn.pack(pady=5)

select_template_btn = Button(root, text="Select Template File", command=select_template_file, bg="#4CAF50", fg="white", font=custom_font)
select_template_btn.pack(pady=5)

Label(root, text="Select the column for Name:", bg="#f0f4f8", font=custom_font).pack(pady=5)
column_menu = OptionMenu(root, selected_column, *column_options)
column_menu.pack(pady=5)

Label(root, text="Select Output Format:", bg="#f0f4f8", font=custom_font).pack(pady=5)
format_menu = OptionMenu(root, output_format, "PDF", "DOCX")
format_menu.pack(pady=5)

generate_btn = Button(root, text="Generate Certificates", command=generate_certificates, bg="#4CAF50", fg="white", font=("Arial", 12, "bold"))
generate_btn.pack(pady=20)

info_label = Label(root, text="Made By Jayesh R Kesari", font=("Arial", 6, "bold"), bg="#4CAF50", fg="white", pady=8)
info_label.pack(fill="x")

# Start the main loop
root.mainloop()
